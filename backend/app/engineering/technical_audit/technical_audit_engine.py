from pathlib import Path
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image
)

from docx import Document
from docx.shared import Inches


TECHNICAL_AUDIT_CHECKLIST = {
    "general_installation": [
        "Availability of updated single-line diagram",
        "Identification and labeling of electrical panels",
        "Availability of electrical maintenance records",
        "Compliance with applicable electrical standards",
        "Accessibility and cleanliness of electrical rooms"
    ],
    "lv_switchboards": [
        "Main LV switchboard rating verification",
        "Short-circuit withstand verification",
        "Busbar condition and tightening verification",
        "Protection device settings verification",
        "Thermal inspection evidence"
    ],
    "protection_coordination": [
        "Upstream/downstream selectivity verification",
        "Breaking capacity verification",
        "Earth fault protection verification",
        "RCD/DDR coordination where applicable",
        "Protection settings documentation"
    ],
    "cables_and_routing": [
        "Cable sizing verification",
        "Voltage drop verification",
        "Cable tray filling verification",
        "Power/control/communication cable separation",
        "Fire resistance and LSZH requirements where applicable"
    ],
    "earthing_and_bonding": [
        "Earthing system identification TT/TN/IT",
        "Earth resistance measurement availability",
        "Equipotential bonding verification",
        "PE/PEN continuity verification",
        "Lightning protection bonding verification"
    ],
    "power_quality": [
        "Power factor verification",
        "Harmonic distortion verification",
        "Neutral overheating risk verification",
        "Capacitor bank resonance risk verification",
        "Need for harmonic filtering"
    ],
    "critical_power": [
        "UPS loading verification",
        "Battery autonomy verification",
        "Generator autonomy verification",
        "ATS/STS transfer logic verification",
        "N+1 or 2N redundancy verification"
    ],
    "safety": [
        "Electrical room safety signage",
        "Arc flash risk screening",
        "Lockout/tagout procedure availability",
        "PPE availability",
        "Emergency shutdown accessibility"
    ]
}


def get_technical_audit_checklist():

    return TECHNICAL_AUDIT_CHECKLIST


def evaluate_technical_audit(
    project_name: str,
    auditor_name: str,
    site_name: str,
    audit_items: list
):

    total = len(audit_items)
    compliant = 0
    non_compliant = 0
    partial = 0
    not_applicable = 0
    critical_findings = []

    for item in audit_items:

        status = item.get("status", "").lower()
        severity = item.get("severity", "medium").lower()

        if status == "compliant":
            compliant += 1
        elif status == "partial":
            partial += 1
        elif status == "not_applicable":
            not_applicable += 1
        else:
            non_compliant += 1

        if status in ["non_compliant", "partial"] and severity == "critical":
            critical_findings.append(item)

    effective_total = max(total - not_applicable, 1)

    score = round(
        (
            compliant + 0.5 * partial
        ) / effective_total * 100,
        2
    )

    if score >= 85 and len(critical_findings) == 0:
        global_status = "ACCEPTABLE"
    elif score >= 65:
        global_status = "ACCEPTABLE WITH RESERVES"
    else:
        global_status = "NON-COMPLIANT"

    recommendations = []

    if critical_findings:
        recommendations.append(
            "Critical findings detected. Immediate corrective actions are required."
        )

    if non_compliant > 0:
        recommendations.append(
            "Non-compliant items must be corrected and verified after remediation."
        )

    if partial > 0:
        recommendations.append(
            "Partially compliant items require engineering review and action plan."
        )

    recommendations.append(
        "Final validation must be performed by a qualified electrical engineer or authorized technical auditor."
    )

    return {
        "project_name": project_name,
        "site_name": site_name,
        "auditor_name": auditor_name,
        "audit_date": datetime.now().strftime("%Y-%m-%d"),
        "summary": {
            "total_items": total,
            "compliant": compliant,
            "partial": partial,
            "non_compliant": non_compliant,
            "not_applicable": not_applicable,
            "score_percent": score,
            "global_status": global_status,
            "critical_findings_count": len(critical_findings)
        },
        "critical_findings": critical_findings,
        "audit_items": audit_items,
        "recommendations": recommendations
    }


def _safe_export_name(name: str):

    return name.replace(" ", "_").replace("/", "_").lower()


def export_audit_pdf(
    audit_data: dict,
    file_name: str,
    logo_1_path: str | None = None,
    logo_2_path: str | None = None
):

    output_dir = Path("app/exports/audits")
    output_dir.mkdir(parents=True, exist_ok=True)

    output_path = output_dir / f"{_safe_export_name(file_name)}.pdf"

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "AuditTitle",
        parent=styles["Title"],
        textColor=colors.HexColor("#0F172A"),
        fontSize=18,
        leading=22,
        spaceAfter=14
    )

    heading_style = ParagraphStyle(
        "AuditHeading",
        parent=styles["Heading2"],
        textColor=colors.HexColor("#1E3A8A"),
        fontSize=13,
        leading=16,
        spaceBefore=12,
        spaceAfter=8
    )

    story = []

    logos = []
    for path in [logo_1_path, logo_2_path]:
        if path and Path(path).exists():
            img = Image(path, width=85, height=45)
            logos.append(img)

    if logos:
        story.append(Table([logos], hAlign="LEFT"))
        story.append(Spacer(1, 10))

    story.append(Paragraph("TECHNICAL ELECTRICAL AUDIT REPORT", title_style))
    story.append(Paragraph(audit_data.get("project_name", ""), styles["Heading3"]))
    story.append(Spacer(1, 8))

    info_table = Table([
        ["Site", audit_data.get("site_name", "")],
        ["Auditor", audit_data.get("auditor_name", "")],
        ["Date", audit_data.get("audit_date", "")],
        ["Global Status", audit_data.get("summary", {}).get("global_status", "")],
        ["Score", f"{audit_data.get('summary', {}).get('score_percent', 0)} %"]
    ], colWidths=[130, 350])

    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E2E8F0")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))

    story.append(info_table)

    story.append(Paragraph("Executive Summary", heading_style))
    summary = audit_data.get("summary", {})
    summary_table = Table([
        ["Total Items", summary.get("total_items", 0)],
        ["Compliant", summary.get("compliant", 0)],
        ["Partial", summary.get("partial", 0)],
        ["Non-Compliant", summary.get("non_compliant", 0)],
        ["Not Applicable", summary.get("not_applicable", 0)],
        ["Critical Findings", summary.get("critical_findings_count", 0)]
    ], colWidths=[180, 280])

    summary_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DBEAFE")),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))

    story.append(summary_table)

    story.append(Paragraph("Audit Findings", heading_style))

    rows = [["Category", "Requirement", "Status", "Severity", "Comment"]]

    for item in audit_data.get("audit_items", []):
        rows.append([
            item.get("category", ""),
            item.get("requirement", ""),
            item.get("status", ""),
            item.get("severity", ""),
            item.get("comment", "")
        ])

    findings_table = Table(rows, repeatRows=1, colWidths=[85, 150, 70, 70, 115])

    findings_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("PADDING", (0, 0), (-1, -1), 4),
    ]))

    story.append(findings_table)

    story.append(Paragraph("Recommendations", heading_style))

    for rec in audit_data.get("recommendations", []):
        story.append(Paragraph(f"- {rec}", styles["BodyText"]))

    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "This document is generated by Electrical AI Platform. Final responsibility remains with the qualified auditor/engineer.",
        styles["Italic"]
    ))

    doc.build(story)

    return {
        "format": "pdf",
        "path": str(output_path),
        "status": "exported"
    }


def export_audit_docx(
    audit_data: dict,
    file_name: str,
    logo_1_path: str | None = None,
    logo_2_path: str | None = None
):

    output_dir = Path("app/exports/audits")
    output_dir.mkdir(parents=True, exist_ok=True)

    output_path = output_dir / f"{_safe_export_name(file_name)}.docx"

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    for path in [logo_1_path, logo_2_path]:
        if path and Path(path).exists():
            document.add_picture(path, width=Inches(1.4))

    document.add_heading("TECHNICAL ELECTRICAL AUDIT REPORT", 0)
    document.add_heading(audit_data.get("project_name", ""), level=1)

    document.add_paragraph(f"Site: {audit_data.get('site_name', '')}")
    document.add_paragraph(f"Auditor: {audit_data.get('auditor_name', '')}")
    document.add_paragraph(f"Date: {audit_data.get('audit_date', '')}")

    summary = audit_data.get("summary", {})

    document.add_heading("Executive Summary", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for key, value in summary.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = str(value)

    document.add_heading("Audit Findings", level=1)
    findings_table = document.add_table(rows=1, cols=5)
    findings_table.style = "Table Grid"

    headers = ["Category", "Requirement", "Status", "Severity", "Comment"]
    for i, header in enumerate(headers):
        findings_table.rows[0].cells[i].text = header

    for item in audit_data.get("audit_items", []):
        row = findings_table.add_row().cells
        row[0].text = item.get("category", "")
        row[1].text = item.get("requirement", "")
        row[2].text = item.get("status", "")
        row[3].text = item.get("severity", "")
        row[4].text = item.get("comment", "")

    document.add_heading("Recommendations", level=1)
    for rec in audit_data.get("recommendations", []):
        document.add_paragraph(rec, style="List Bullet")

    document.add_paragraph(
        "Final validation must be performed by a qualified electrical engineer or authorized technical auditor."
    )

    document.save(str(output_path))

    return {
        "format": "docx",
        "path": str(output_path),
        "status": "exported"
    }


def export_technical_audit(
    audit_data: dict,
    file_name: str,
    export_format: str,
    include_logo_1: bool = False,
    logo_1_path: str | None = None,
    include_logo_2: bool = False,
    logo_2_path: str | None = None
):

    selected_logo_1 = logo_1_path if include_logo_1 else None
    selected_logo_2 = logo_2_path if include_logo_2 else None

    if export_format.lower() == "docx":
        return export_audit_docx(
            audit_data=audit_data,
            file_name=file_name,
            logo_1_path=selected_logo_1,
            logo_2_path=selected_logo_2
        )

    return export_audit_pdf(
        audit_data=audit_data,
        file_name=file_name,
        logo_1_path=selected_logo_1,
        logo_2_path=selected_logo_2
    )
