from pathlib import Path
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    Image
)

from docx import Document
from docx.shared import Inches


def _safe_name(name: str):

    return (
        name.replace(" ", "_")
        .replace("/", "_")
        .lower()
    )


def _build_cover_page(
    story,
    report_title,
    project_name,
    company_name,
    logos
):

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CoverTitle",
        parent=styles["Title"],
        fontSize=24,
        leading=30,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=24
    )

    subtitle_style = ParagraphStyle(
        "CoverSubtitle",
        parent=styles["Heading2"],
        fontSize=15,
        leading=20,
        textColor=colors.HexColor("#334155"),
        spaceAfter=14
    )

    if logos:

        logo_images = []

        for path in logos:

            if Path(path).exists():

                logo_images.append(
                    Image(
                        path,
                        width=120,
                        height=60
                    )
                )

        if logo_images:

            story.append(
                Table([logo_images])
            )

            story.append(
                Spacer(1, 24)
            )

    story.append(
        Paragraph(
            report_title,
            title_style
        )
    )

    story.append(
        Paragraph(
            project_name,
            subtitle_style
        )
    )

    story.append(
        Paragraph(
            company_name,
            subtitle_style
        )
    )

    story.append(
        Spacer(1, 30)
    )

    story.append(
        Paragraph(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            styles["BodyText"]
        )
    )

    story.append(
        Spacer(1, 450)
    )

    story.append(
        Paragraph(
            "Electrical AI Platform - Engineering Report System",
            styles["Italic"]
        )
    )

    story.append(PageBreak())


def export_premium_pdf(
    file_name: str,
    report_title: str,
    project_name: str,
    company_name: str,
    sections: list,
    logos: list | None = None
):

    output_dir = Path(
        "app/exports/premium_reports"
    )

    output_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    output_path = output_dir / f"{_safe_name(file_name)}.pdf"

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=32,
        leftMargin=32,
        topMargin=32,
        bottomMargin=32
    )

    styles = getSampleStyleSheet()

    heading_style = ParagraphStyle(
        "HeadingStyle",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#1E3A8A"),
        spaceAfter=10,
        spaceBefore=14
    )

    body_style = ParagraphStyle(
        "BodyStyle",
        parent=styles["BodyText"],
        fontSize=10,
        leading=15
    )

    story = []

    _build_cover_page(
        story=story,
        report_title=report_title,
        project_name=project_name,
        company_name=company_name,
        logos=logos or []
    )

    for section in sections:

        story.append(
            Paragraph(
                section.get("title", ""),
                heading_style
            )
        )

        story.append(
            Spacer(1, 6)
        )

        content = section.get("content", "")

        if isinstance(content, list):

            rows = [["Item", "Value"]]

            for item in content:

                rows.append([
                    str(item.get("label", "")),
                    str(item.get("value", ""))
                ])

            table = Table(
                rows,
                colWidths=[220, 280]
            )

            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]))

            story.append(table)

        else:

            story.append(
                Paragraph(
                    str(content),
                    body_style
                )
            )

        story.append(
            Spacer(1, 14)
        )

    story.append(
        Spacer(1, 24)
    )

    story.append(
        Paragraph(
            "Final engineering validation remains under responsibility of qualified personnel.",
            styles["Italic"]
        )
    )

    doc.build(story)

    return {
        "format": "pdf",
        "path": str(output_path),
        "status": "exported"
    }


def export_premium_docx(
    file_name: str,
    report_title: str,
    project_name: str,
    company_name: str,
    sections: list,
    logos: list | None = None
):

    output_dir = Path(
        "app/exports/premium_reports"
    )

    output_dir.mkdir(
        parents=True,
        exist_ok=True
    )

    output_path = output_dir / f"{_safe_name(file_name)}.docx"

    document = Document()

    section = document.sections[0]

    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    if logos:

        for logo in logos:

            if Path(logo).exists():

                document.add_picture(
                    logo,
                    width=Inches(1.7)
                )

    document.add_heading(
        report_title,
        level=0
    )

    document.add_heading(
        project_name,
        level=1
    )

    document.add_paragraph(
        company_name
    )

    document.add_paragraph(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )

    for section_data in sections:

        document.add_heading(
            section_data.get("title", ""),
            level=1
        )

        content = section_data.get("content", "")

        if isinstance(content, list):

            table = document.add_table(
                rows=1,
                cols=2
            )

            table.style = "Table Grid"

            hdr = table.rows[0].cells

            hdr[0].text = "Item"
            hdr[1].text = "Value"

            for item in content:

                row = table.add_row().cells

                row[0].text = str(
                    item.get("label", "")
                )

                row[1].text = str(
                    item.get("value", "")
                )

        else:

            document.add_paragraph(
                str(content)
            )

    document.add_paragraph(
        "Final engineering validation remains under responsibility of qualified personnel."
    )

    document.save(
        str(output_path)
    )

    return {
        "format": "docx",
        "path": str(output_path),
        "status": "exported"
    }


def export_premium_report(
    file_name: str,
    report_title: str,
    project_name: str,
    company_name: str,
    sections: list,
    export_format: str = "pdf",
    logos: list | None = None
):

    if export_format.lower() == "docx":

        return export_premium_docx(
            file_name=file_name,
            report_title=report_title,
            project_name=project_name,
            company_name=company_name,
            sections=sections,
            logos=logos
        )

    return export_premium_pdf(
        file_name=file_name,
        report_title=report_title,
        project_name=project_name,
        company_name=company_name,
        sections=sections,
        logos=logos
    )
